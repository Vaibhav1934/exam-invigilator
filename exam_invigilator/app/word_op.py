from docx import Document
from docx.shared import Inches,Pt
import os 
from docx.shared import Mm
from docx.enum.table import WD_TABLE_ALIGNMENT
import aspose.words as aw

def pdf_conv(ALL,data,pos):


	
	print("yes")
	if ALL:
		print("yes")
		document = Document()
		section = document.sections[0]
		section.page_height = Mm(297)
		section.page_width = Mm(210)
		section.left_margin = Mm(25.4)
		section.right_margin = Mm(25.4)
		section.top_margin = Mm(25.4)
		section.bottom_margin = Mm(25.4)
		section.header_distance = Mm(12.7)
		section.footer_distance = Mm(12.7)
		
		paragraph=document.add_paragraph()
		paragraph_format=paragraph.paragraph_format
	
		
		paragraph_format.right_indent = Pt(24)
		
		
		
		
		
		document.add_heading('SRINIVAS INSTITUTE OF TECHNOLOGY , MANGALURU', 0)
		
		style = document.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = Pt(14)	
		document.add_paragraph("			")
		para3=document.add_paragraph('Following members are required to take up the university theory examination duty on following days')
		para3.paragraph_format.left_indent = Pt(48)
		document.add_paragraph("			")
		table = document.add_table(rows=1, cols=3)
		table.autofit = False 
		table.allow_autofit = False
		hdr_cells = table.rows[0].cells
		date=hdr_cells[0].text = 'Date'
		date=hdr_cells[0].paragraphs[0].runs[0]
		date.font.bold=True
		date=hdr_cells[1].text = 'Name'
		date=hdr_cells[1].paragraphs[0].runs[0]
		date.font.bold=True
		date=hdr_cells[2].text = 'session'
		date=hdr_cells[2].paragraphs[0].runs[0]
		date.font.bold=True
		for date,Name,ses  in data:
		    row_cells = table.add_row().cells
		    row_cells[0].text = str(date)
		    row_cells[1].text = Name
		    if ses=='M':
		    	row_cells[2].text = 'Morning'
		    else:
		    	row_cells[2].text = 'Afternoon'
		    
		    table.columns[0].width = Inches(1.5)
		    table.rows[0].cells[0].width = Inches(1.5)
		table.alignment = WD_TABLE_ALIGNMENT.CENTER

		sign=document.add_paragraph("PRINCIPAL & CHEIF SUPERINTENDENT")
		sign.paragraph_format.left_indent = Pt(48)
		sign.paragraph_format.space_before = Inches(1.0)
		p=sign.add_run()
		p.font.size = Pt(1)
		document.add_page_break()
	
	
		document.save('duty_list.docx')
		doc = aw.Document("duty_list.docx")
		doc.save("duty_list.pdf")
	
		print(os.getcwd())
	

	else:
	
	
			print("No")
			document = Document()
			section = document.sections[0]
			section.page_height = Mm(297)
			section.page_width = Mm(210)
			section.left_margin = Mm(25.4)
			section.right_margin = Mm(25.4)
			section.top_margin = Mm(25.4)
			section.bottom_margin = Mm(25.4)
			section.header_distance = Mm(12.7)
			section.footer_distance = Mm(12.7)		
			for date,Name,session  in data:
				print(data)
				print(Name,session)
				a=str(date)

		
				paragraph=document.add_paragraph()
				paragraph_format=paragraph.paragraph_format
	
		
				paragraph_format.right_indent = Pt(24)
		
			
					
	
		
				document.add_heading('SRINIVAS INSTITUTE OF TECHNOLOGY , MANGALURU', 0)
		
				style = document.styles['Normal']
				font = style.font
				font.name = 'Times New Roman'
				font.size = Pt(14)
				para=document.add_paragraph('To: ')
				para.paragraph_format.left_indent = Pt(24)
				para.bold=True
				para2=document.add_paragraph(Name)
				para2.paragraph_format.left_indent = Pt(48)
				para.paragraph_format.space_before = Inches(0.25)
				para3=document.add_paragraph('You are required to take up the university theory examination duty on following days')
				para3.paragraph_format.left_indent = Pt(48)
				document.add_paragraph("			")
				table = document.add_table(rows=1, cols=3)
				table.autofit = False 
				table.allow_autofit = False
				hdr_cells = table.rows[0].cells
				date=hdr_cells[0].text = 'Date'
				date=hdr_cells[0].paragraphs[0].runs[0]
				date.font.bold=True
				date=hdr_cells[1].text = 'Name'
				date=hdr_cells[1].paragraphs[0].runs[0]
				date.font.bold=True
				date=hdr_cells[2].text = 'session'
				date=hdr_cells[2].paragraphs[0].runs[0]
				date.font.bold=True
				
				row_cells = table.add_row().cells

				row_cells[0].text =a
			
				row_cells[1].text = Name
				if session=='M':
				    	row_cells[2].text = 'Morning'
				else:
				    	row_cells[2].text = 'Afternoon'
				table.columns[0].width = Inches(1.5)
				table.rows[0].cells[0].width = Inches(1.5)
				table.alignment = WD_TABLE_ALIGNMENT.CENTER
		
				sign=document.add_paragraph("PRINCIPAL & CHEIF SUPERINTENDENT")
				sign.paragraph_format.left_indent = Pt(48)
				sign.paragraph_format.space_before = Inches(1.0)
				p=sign.add_run()
				p.font.size = Pt(1)
				document.add_page_break()
		
			if pos=="inv":
				document.save('inv.docx')
				doc = aw.Document('inv.docx')
				doc.save("inv.pdf")
			else:
				document.save('dc.docx')
				doc = aw.Document('dc.docx')
				doc.save("dc.pdf")			
				
			#os.
			
			print(os.getcwd())	
	
		
	






